from __future__ import annotations

import math
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT_DIR = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT_DIR / "docs"
OUTPUT_PATH = DOCS_DIR / "Tai_lieu_tong_hop_ky_thuat_Cherry_Core_2026.docx"
SOURCE_PLAN_PATH = Path(
    r"E:\Freelance\Research\D12_02.2026_NCKH2026\KH thuc hien SVNCKH 2026.docx"
)

STATUS_LEGEND = [
    ("đã có code", "Mã nguồn hoặc adapter đã hiện diện trong lõi hệ thống."),
    ("đã có script kiểm thử", "Có script kiểm thử hoặc probe xác minh riêng."),
    ("đã có output minh chứng", "Có artifact đầu ra hoặc báo cáo chạy thử trong repo."),
    ("mới ở mức adapter/probe", "Mới dừng ở adapter, tài liệu nghiên cứu hoặc bước chuẩn bị."),
]

ABBREVIATIONS = [
    ("ASR", "Automatic Speech Recognition - nhận dạng tiếng nói tự động"),
    ("VAD", "Voice Activity Detection - phát hiện đoạn có tiếng nói"),
    ("LLM", "Large Language Model - mô hình ngôn ngữ lớn"),
    ("WER", "Word Error Rate - tỷ lệ lỗi từ trong ASR"),
    ("DER", "Diarization Error Rate - tỷ lệ lỗi phân định người nói"),
    ("JSON", "JavaScript Object Notation - định dạng dữ liệu có cấu trúc"),
    ("VBx", "Variational Bayes / Viterbi-based resegmentation cho diarization"),
    ("CoVe", "Chain-of-Verification - chuỗi xác minh nhiều bước cho LLM"),
    ("GBNF", "Grammar-Based Normal Form dùng để ràng buộc sinh văn bản"),
    ("DTO", "Data Transfer Object - đối tượng truyền dữ liệu giữa các lớp"),
    ("CLI", "Command Line Interface - giao diện dòng lệnh"),
    ("GUI", "Graphical User Interface - giao diện đồ họa người dùng"),
]

GLOSSARY = [
    (
        "Hallucination",
        "Hiện tượng mô hình sinh ra văn bản không tồn tại trong âm thanh hoặc không được chứng minh bởi dữ liệu đầu vào.",
    ),
    ("Anti-hallucination", "Tập hợp kỹ thuật cấu hình, tiền xử lý và hậu xử lý để giảm nội dung bịa."),
    ("Diarization", "Quá trình xác định ai nói câu nào theo trục thời gian."),
    ("Alignment", "Quá trình gắn mốc thời gian cho từ hoặc câu để liên kết ASR với speaker segment."),
    (
        "Offline-first",
        "Nguyên tắc thiết kế cho phép hệ thống tiếp tục hoạt động sau khi đã cài đặt model cục bộ, không phụ thuộc internet.",
    ),
    ("Forensic report", "Báo cáo trích xuất thông tin điều tra, có cấu trúc, phục vụ phân tích nghiệp vụ."),
    ("Artifact", "Tệp đầu ra, log, transcript, báo cáo, hoặc metadata sinh ra trong một job xử lý."),
]

REFERENCES = [
    "[1] Kế hoạch thực hiện nhiệm vụ nghiên cứu khoa học học viên năm học 2025-2026, file nguồn do người dùng cung cấp.",
    "[2] E:/research/Cherry2/cherry_core/docs/DEEP_RESEARCH_UPGRADE_2026.md - báo cáo nghiên cứu chuyên sâu về anti-hallucination, diarization và structured output.",
    "[3] E:/research/Cherry2/cherry_core/docs/SOTA_DIARIZATION_RESEARCH.md - khảo sát hướng diarization SOTA và VBx resegmentation.",
    "[4] E:/research/Cherry2/cherry_core/docs/SCIENTIFIC_PROMPT_RESEARCH.md - khung prompt forensic dựa trên SVA/SCAN/ACH.",
    "[5] E:/research/Cherry2/cherry_core/docs/WHISPER_MODEL_SELECTION.md - cơ sở chọn Whisper V2 làm baseline ít hallucination hơn V3.",
    "[6] E:/research/Cherry2/cherry_core/docs/CYCLE_4_DESIGN_VLLM.md và CYCLE_4_STATUS.md - thiết kế backend LLM hiệu năng cao và cơ chế fallback.",
    "[7] E:/research/Cherry2/cherry_core/PROOF_RESULTS.md - bằng chứng chạy thử correction, anti-hallucination và GPU.",
]


def set_run_font(run, size: float = 13, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(13)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True


def add_center_paragraph(document: Document, text: str, size: float = 13, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_justified_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_run_font(run)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(item)
        set_run_font(run)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size=11, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(str(value))
            set_run_font(run, size=11)


def add_page_break(document: Document) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def human_size(num_bytes: int) -> str:
    if num_bytes <= 0:
        return "0 B"
    units = ["B", "KB", "MB", "GB", "TB"]
    power = min(int(math.log(num_bytes, 1024)), len(units) - 1)
    value = num_bytes / (1024**power)
    return f"{value:.2f} {units[power]}"


def safe_relative(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT_DIR)).replace("\\", "/")
    except ValueError:
        return str(path)


TECHNIQUE_STATUS = [
    {
        "name": "Whisper V2 với cấu hình chống ảo giác",
        "status": "đã có output minh chứng",
        "summary": "ASR nền hiện tại của Cherry Core, dùng cấu hình giải mã bảo thủ và bộ lọc lặp hậu xử lý.",
        "evidence": [
            "infrastructure/adapters/asr/whisperv2_adapter.py",
            "PROOF_RESULTS.md",
            "output/transcript_v2_tuned.txt",
        ],
    },
    {
        "name": "PhoWhisper cho ASR tiếng Việt",
        "status": "đã có code",
        "summary": "Adapter riêng cho PhoWhisper với model offline cục bộ, ưu tiên nhánh safetensors.",
        "evidence": [
            "infrastructure/adapters/asr/phowhisper_adapter.py",
            "models/phowhisper-safe/",
            "scripts/setup_models.py",
        ],
    },
    {
        "name": "WhisperX end-to-end (ASR + alignment + diarization)",
        "status": "đã có script kiểm thử",
        "summary": "Nhánh nâng cấp tích hợp faster-whisper, word alignment và pyannote; đã có script setup và verify.",
        "evidence": [
            "infrastructure/adapters/asr/whisperx_adapter.py",
            "scripts/setup_whisperx_offline.py",
            "scripts/verify_whisperx.py",
        ],
    },
    {
        "name": "Decoding thresholds giảm hallucination",
        "status": "đã có code",
        "summary": "Các ngưỡng giải mã được hard-code trong WhisperV2Adapter để giảm nội dung bịa ở vùng im lặng.",
        "evidence": ["infrastructure/adapters/asr/whisperv2_adapter.py"],
    },
    {
        "name": "Silero VAD tiền xử lý",
        "status": "đã có code",
        "summary": "VAD chạy offline bằng model JIT, ưu tiên bảo toàn thông tin với ngưỡng bảo thủ.",
        "evidence": [
            "infrastructure/adapters/vad/silero_adapter.py",
            "models/silero/",
        ],
    },
    {
        "name": "Bag-of-Hallucinations tiếng Việt và tiếng Anh",
        "status": "đã có code",
        "summary": "Tập biểu thức gây ảo giác thường gặp được mã hóa trực tiếp trong HallucinationFilter.",
        "evidence": ["infrastructure/adapters/asr/hallucination_filter.py"],
    },
    {
        "name": "Delooping hậu xử lý",
        "status": "đã có code",
        "summary": "Lớp khử lặp theo từ và theo cụm nhằm loại bỏ vòng lặp nhiều lần trong bản chép lời.",
        "evidence": [
            "infrastructure/adapters/asr/hallucination_filter.py",
            "infrastructure/adapters/asr/whisperv2_adapter.py",
        ],
    },
    {
        "name": "Pyannote diarization",
        "status": "mới ở mức adapter/probe",
        "summary": "Adapter hỗ trợ Pyannote Community-1/3.1 nhưng vẫn phụ thuộc token Hugging Face và chưa có benchmark rõ trong repo.",
        "evidence": [
            "infrastructure/adapters/diarization/pyannote_adapter.py",
            "docs/CYCLE_2_RESEARCH_DIARIZATION.md",
        ],
    },
    {
        "name": "SpeechBrain ECAPA-TDNN diarization",
        "status": "đã có code",
        "summary": "Nhánh diarization offline dựa trên embedding ECAPA-TDNN, clustering phổ và khối hợp nhất segment.",
        "evidence": [
            "infrastructure/adapters/diarization/speechbrain_adapter.py",
            "models/speechbrain/",
        ],
    },
    {
        "name": "VBx/Viterbi resegmentation",
        "status": "đã có code",
        "summary": "Bộ làm mượt nhãn speaker dựa trên HMM/Viterbi để giảm speaker switching ngắn hạn.",
        "evidence": [
            "infrastructure/adapters/diarization/vbx_refiner.py",
            "infrastructure/adapters/diarization/speechbrain_adapter.py",
        ],
    },
    {
        "name": "Word-speaker alignment bằng IntervalTree",
        "status": "đã có code",
        "summary": "Dịch vụ căn chỉnh timestamp word-level của ASR sang segment người nói, hỗ trợ tra cứu nhanh bằng IntervalTree.",
        "evidence": ["core/services/alignment_service.py"],
    },
    {
        "name": "Rule-based phonetic correction",
        "status": "đã có code",
        "summary": "Bộ hiệu chỉnh lỗi đồng âm dị nghĩa tiếng Việt bằng từ điển JSON và biểu thức chính quy.",
        "evidence": [
            "application/services/phonetic_corrector.py",
            "assets/vocab/vietnamese_phonetic_errors.json",
        ],
    },
    {
        "name": "ProtonX Seq2Seq correction",
        "status": "đã có output minh chứng",
        "summary": "Adapter hậu hiệu chỉnh câu chữ bằng mô hình ProtonX chạy local, hỗ trợ chia đoạn transcript dài.",
        "evidence": [
            "infrastructure/adapters/correction/protonx_adapter.py",
            "PROOF_RESULTS.md",
            "models/protonx/",
        ],
    },
    {
        "name": "LLM contextual correction",
        "status": "đã có code",
        "summary": "CorrectionService kết hợp rule-based, vocabulary injection và LLM cục bộ để hiệu chỉnh theo ngữ cảnh miền.",
        "evidence": [
            "application/services/correction_service.py",
            "application/services/phonetic_corrector.py",
        ],
    },
    {
        "name": "Prompt modules và scenario YAML",
        "status": "đã có script kiểm thử",
        "summary": "Prompt điều tra được xây theo mô-đun Jinja2 và dữ liệu kịch bản YAML, đã có test render và test deep analysis.",
        "evidence": [
            "application/services/prompt_manager.py",
            "tests/test_prompt_renderer.py",
            "tests/test_deep_analysis.py",
        ],
    },
    {
        "name": "JSON-only prompting và GBNF",
        "status": "đã có code",
        "summary": "Prompt deep investigation ép đầu ra JSON và adapter llama.cpp có chỗ móc grammar, nhưng sample output cho thấy JSON chưa ổn định tuyệt đối.",
        "evidence": [
            "prompts/templates/deep_investigation.j2",
            "prompts/grammars/json_schema.gbnf",
            "application/services/analysis_service.py",
        ],
    },
    {
        "name": "LLM offline qua llama.cpp",
        "status": "đã có output minh chứng",
        "summary": "LlamaCppAdapter là backend LLM mặc định, dùng model GGUF và đã tạo được báo cáo phân tích mẫu trong output.",
        "evidence": [
            "infrastructure/adapters/llm/llamacpp_adapter.py",
            "output/Tiếp nhận yêu cầu đặt phòng của khách lẻ qua điện thoại_analysis_report.json",
        ],
    },
    {
        "name": "vLLM fallback qua WSL2/probe",
        "status": "đã có script kiểm thử",
        "summary": "Adapter hiệu năng cao đã có, kèm script chuẩn bị WSL2 và tài liệu chu trình fallback về llama.cpp nếu Windows không phù hợp.",
        "evidence": [
            "infrastructure/adapters/llm/vllm_adapter.py",
            "scripts/setup_wsl_vllm.ps1",
            "docs/CYCLE_4_STATUS.md",
        ],
    },
]

PROPOSALS = [
    {
        "name": "Calm-Whisper",
        "benefit": "Giảm mạnh hallucination của Whisper bằng tinh chỉnh attention head gây lỗi, giữ WER gần như ổn định.",
        "gap": "Cherry Core mới dừng ở cấu hình giải mã, VAD và BoH; chưa có quy trình fine-tune hoặc mô hình Calm-Whisper nội bộ.",
        "risk": "Cần tài nguyên huấn luyện, tập dữ liệu nhiễu phù hợp và quy trình đánh giá riêng cho tiếng Việt.",
        "sources": ["docs/DEEP_RESEARCH_UPGRADE_2026.md", "docs/UPGRADE_RESEARCH_2026.md"],
    },
    {
        "name": "Chain-of-Verification (CoVe)",
        "benefit": "Tăng độ tin cậy của LLM forensic bằng chu trình nháp - kiểm chứng - tổng hợp nhiều bước.",
        "gap": "Repo hiện có phân tích LLM một lượt; chưa có service tách riêng các vòng xác minh hoặc scoring factual consistency.",
        "risk": "Độ trễ tăng đáng kể, kéo theo nhu cầu cache và thiết kế job nền nếu muốn đưa vào GUI.",
        "sources": ["docs/DEEP_RESEARCH_UPGRADE_2026.md", "docs/ADVANCED_RESEARCH_PROPOSAL.md"],
    },
    {
        "name": "XGrammar / structured decoding",
        "benefit": "Nâng valid rate của JSON đầu ra lên mức gần tuyệt đối, phù hợp báo cáo chứng cứ có cấu trúc.",
        "gap": "Cherry Core mới có GBNF cho llama.cpp và prompt JSON-only; chưa tích hợp XGrammar hoặc guided decoding thực sự.",
        "risk": "Phụ thuộc backend LLM, có thể đòi hỏi chuyển hạ tầng sang vLLM hoặc một gateway riêng.",
        "sources": ["docs/DEEP_RESEARCH_UPGRADE_2026.md", "docs/CYCLE_4_DESIGN_VLLM.md"],
    },
    {
        "name": "Semantic retrieval thay keyword matching",
        "benefit": "Lấy vocabulary và tri thức miền theo ngữ nghĩa thay vì trigger từ khóa thô, giúp correction bám sát ngữ cảnh hơn.",
        "gap": "CorrectionService hiện đang chọn từ vựng bằng danh sách trigger thủ công; chưa có FAISS hoặc sentence-transformers.",
        "risk": "Cần thêm pipeline index offline, đánh giá chi phí bộ nhớ và cơ chế cập nhật tri thức.",
        "sources": ["docs/DEEP_RESEARCH_UPGRADE_2026.md"],
    },
    {
        "name": "Pyannote Community-1 / 4.0",
        "benefit": "Cải thiện speaker assignment, speaker counting và hỗ trợ overlap tự nhiên hơn so với nhánh cũ.",
        "gap": "Adapter đã nhắc tới Community-1 nhưng repo chưa có benchmark DER hay artifact xác thực chính thức.",
        "risk": "Token, cache model, license và tương thích thiết bị là các điểm triển khai cần quản trị.",
        "sources": ["infrastructure/adapters/diarization/pyannote_adapter.py", "docs/SOTA_DIARIZATION_RESEARCH.md"],
    },
    {
        "name": "WhisperX migration",
        "benefit": "Hợp nhất ASR, alignment và diarization theo một pipeline end-to-end, phù hợp GUI timeline người nói.",
        "gap": "Handoff cho thấy migration đang ở trạng thái kiểm chứng offline; pipeline mặc định của hệ chưa chuyển hẳn sang WhisperX.",
        "risk": "Model nặng, phụ thuộc alignment/diarization assets và cần benchmark so với nhánh legacy.",
        "sources": ["AGENT_HANDOFF.md", "infrastructure/adapters/asr/whisperx_adapter.py"],
    },
    {
        "name": "vLLM qua WSL2 hoặc remote serving",
        "benefit": "Tăng throughput, hỗ trợ batching liên tục và mở đường cho guided decoding cấu trúc.",
        "gap": "Adapter và tài liệu đã có, nhưng hiện trạng vẫn lấy llama.cpp làm fallback an toàn; chưa có mặt bằng triển khai ổn định cho GUI.",
        "risk": "Windows native còn nhạy cảm; nếu tách thành service riêng thì phải bổ sung cơ chế health check và audit truy cập.",
        "sources": ["docs/CYCLE_4_DESIGN_VLLM.md", "docs/CYCLE_4_STATUS.md"],
    },
]

GUI_SCREENS = [
    {
        "name": "Tổng quan hồ sơ",
        "goal": "Theo dõi danh sách ca xử lý, trạng thái, mức độ ưu tiên và đường dẫn artifact.",
        "inputs": "Bộ lọc hồ sơ, trạng thái job, ngày xử lý, scenario phân tích.",
        "outputs": "Danh sách job, đường dẫn transcript, báo cáo, nhật ký xử lý.",
        "core_mapping": "Bọc file-based artifact trong output/ và metadata job cục bộ.",
    },
    {
        "name": "Tạo phiên xử lý audio",
        "goal": "Nạp file âm thanh, chọn pipeline, chọn model và bật các tùy chọn correction/diarization.",
        "inputs": "Audio file, ASR engine, bật hoặc tắt diarization, bật correction, scenario.",
        "outputs": "Job mới với tham số chạy và đường dẫn input.",
        "core_mapping": "Gọi TranscribeAudioUseCase, SystemFactory, adapter ASR, diarizer và corrector.",
    },
    {
        "name": "Theo dõi job",
        "goal": "Hiển thị tiến độ chạy, cảnh báo lỗi, thông tin model, thời gian xử lý và log.",
        "inputs": "Job id, websocket stream hoặc polling.",
        "outputs": "Tiến độ, trạng thái, lỗi, dấu mốc pipeline.",
        "core_mapping": "Wrapper quanh orchestration layer mới; dùng log từ adapter và job metadata.",
    },
    {
        "name": "Transcript thô và transcript đã hiệu chỉnh",
        "goal": "So sánh bản raw, corrected, đánh dấu chỗ bị nghi ngờ sai hoặc có thể chỉnh tay.",
        "inputs": "Transcript id, bộ lọc segment, thao tác chỉnh tay.",
        "outputs": "Transcript hiện hành, diff raw-corrected, lịch sử chỉnh sửa.",
        "core_mapping": "Dùng Transcript, CorrectionService, ProtonXAdapter, phonetic_corrector.",
    },
    {
        "name": "Timeline người nói",
        "goal": "Xem segment theo thời gian, speaker block, word alignment và đoạn overlap nghi vấn.",
        "inputs": "Transcript segment, speaker segment, waveform hoặc timeline navigation.",
        "outputs": "Speaker-aware transcript và nhãn speaker theo đoạn.",
        "core_mapping": "Dùng SpeakerSegment, AlignmentService, adapter diarization và WhisperX.",
    },
    {
        "name": "Panel phân tích LLM",
        "goal": "Chạy phân tích forensic, xem JSON hoặc raw output, entities, timeline và mức đe dọa.",
        "inputs": "Transcript đã duyệt, scenario, prompt profile.",
        "outputs": "Analysis report, extracted entities, hành vi, khuyến nghị.",
        "core_mapping": "Dùng AnalysisService, PromptManager, LlamaCppAdapter, vLLMAdapter.",
    },
    {
        "name": "Xuất báo cáo và chứng cứ",
        "goal": "Đóng gói transcript, JSON báo cáo, nhật ký xử lý và bản DOCX nghiệp vụ.",
        "inputs": "Report id, định dạng xuất, danh sách artifact cần gộp.",
        "outputs": "Gói xuất bản TXT, JSON, DOCX kèm metadata.",
        "core_mapping": "Dùng OutputFormatter, serializer báo cáo, artifact writer mới.",
    },
    {
        "name": "Kiểm tra model và offline assets",
        "goal": "Kiểm tra model có sẵn, dung lượng, trạng thái nạp, token yêu cầu và khả năng chạy offline.",
        "inputs": "Lệnh refresh inventory, yêu cầu preload model.",
        "outputs": "Danh sách model, kích thước, engine, trạng thái load, cảnh báo phụ thuộc.",
        "core_mapping": "Dùng core/config.py, SystemFactory, thư mục models/.",
    },
]

DTO_DEFINITIONS = {
    "JobStatus": [
        ("job_id", "string", "Mã định danh job cục bộ, duy nhất theo phiên xử lý."),
        ("state", "queued|running|completed|failed|cancelled", "Trạng thái xử lý hiện thời."),
        ("stage", "string", "Tên bước đang chạy: upload, transcribe, diarize, correct, analyze, export."),
        ("progress", "number", "Tiến độ 0-100 cho UI hiển thị thanh trạng thái."),
        ("created_at", "datetime", "Thời điểm tạo job."),
        ("updated_at", "datetime", "Thời điểm cập nhật gần nhất."),
        ("error", "string|null", "Thông báo lỗi cuối nếu job thất bại."),
        ("artifacts", "string[]", "Danh sách đường dẫn artifact sinh ra."),
    ],
    "TranscriptView": [
        ("transcript_id", "string", "Mã transcript gắn với một job."),
        ("job_id", "string", "Job nguồn đã sinh transcript."),
        ("language", "string", "Ngôn ngữ xử lý, mặc định vi."),
        ("raw_text", "string", "Bản transcript thô sau ASR."),
        ("corrected_text", "string", "Bản transcript sau correction hoặc chỉnh tay."),
        ("segments", "TranscriptSegmentView[]", "Danh sách segment theo thời gian."),
        ("speakers", "SpeakerSegmentView[]", "Danh sách segment người nói."),
        ("metadata", "object", "Thông tin model, flags anti-hallucination, output paths."),
    ],
    "TranscriptSegmentView": [
        ("segment_id", "string", "Định danh segment."),
        ("start", "number", "Thời điểm bắt đầu theo giây."),
        ("end", "number", "Thời điểm kết thúc theo giây."),
        ("text", "string", "Nội dung segment."),
        ("speaker_id", "string|null", "Speaker gán cho segment, nếu có."),
        ("confidence", "number|null", "Độ tin cậy nếu backend cung cấp."),
        ("words", "object[]", "Danh sách từ có timestamp chi tiết."),
    ],
    "SpeakerSegmentView": [
        ("speaker_id", "string", "Nhãn speaker chuẩn hóa, ví dụ SPEAKER_1."),
        ("start_time", "number", "Thời điểm bắt đầu."),
        ("end_time", "number", "Thời điểm kết thúc."),
        ("text", "string", "Văn bản gán cho speaker sau alignment."),
        ("words", "object[]", "Danh sách word-level alignment nếu có."),
    ],
    "AnalysisReportView": [
        ("report_id", "string", "Mã báo cáo."),
        ("job_id", "string", "Liên kết job gốc."),
        ("scenario", "string", "Kịch bản điều tra được dùng."),
        ("summary", "string", "Tóm tắt cấp hồ sơ."),
        ("threat_level", "string", "Mức đe dọa chuẩn hóa."),
        ("classification", "string", "Phân loại hội thoại."),
        ("structured_data", "object", "Nội dung JSON có cấu trúc."),
        ("raw_output", "string|null", "Đầu ra LLM nguyên bản để đối chiếu."),
        ("validation_status", "string", "Trạng thái kiểm tra schema hoặc JSON."),
    ],
    "ModelInfo": [
        ("model_id", "string", "Tên logic của model hoặc engine."),
        ("family", "asr|diarization|correction|llm|vad", "Nhóm chức năng."),
        ("path", "string", "Đường dẫn cục bộ tới model."),
        ("size_gb", "number", "Dung lượng ước tính."),
        ("offline_ready", "boolean", "Đã có file local đầy đủ hay chưa."),
        ("load_status", "unloaded|loaded|error", "Trạng thái nạp hiện tại."),
        ("notes", "string", "Ràng buộc đặc biệt như HF token, CUDA, WSL2."),
    ],
    "AuditEntry": [
        ("entry_id", "string", "Mã nhật ký."),
        ("job_id", "string|null", "Job liên quan."),
        ("timestamp", "datetime", "Thời điểm phát sinh sự kiện."),
        ("actor", "system|user", "Nguồn phát sinh."),
        ("action", "string", "Hành động: upload, transcribe, manual_edit, export."),
        ("details", "object", "Thông tin bổ sung phục vụ truy vết."),
    ],
}

ENDPOINT_SPECS = [
    {
        "method": "POST",
        "path": "/jobs/transcribe",
        "purpose": "Tạo job phiên âm mới từ audio đầu vào.",
        "request": "Multipart upload gồm file audio, engine ASR, cờ diarization, cờ correction và scenario.",
        "response": "Trả về JobStatus ban đầu cùng đường dẫn tới transcript khi hoàn thành.",
        "errors": "400 file lỗi; 409 model chưa sẵn; 500 lỗi adapter hoặc thiếu tài nguyên GPU/CPU.",
        "mapping": "Wrapper mới quanh TranscribeAudioUseCase, SystemFactory, adapter ASR, VAD, diarizer và corrector.",
    },
    {
        "method": "POST",
        "path": "/jobs/analyze",
        "purpose": "Khởi động pha phân tích forensic cho transcript đã chọn.",
        "request": "JSON gồm transcript_id, scenario, prompt_profile, cờ lưu raw output.",
        "response": "Trả về JobStatus của pha phân tích và report_id khi hoàn thành.",
        "errors": "404 transcript không tồn tại; 409 transcript chưa sẵn sàng; 500 lỗi LLM.",
        "mapping": "Wrapper mới quanh AnalysisService, PromptManager, LlamaCppAdapter hoặc vLLMAdapter.",
    },
    {
        "method": "GET",
        "path": "/jobs/{id}",
        "purpose": "Lấy trạng thái và tiến độ thực thi của một job.",
        "request": "Path parameter id.",
        "response": "Một đối tượng JobStatus cập nhật mới nhất.",
        "errors": "404 job không tồn tại.",
        "mapping": "Lớp read-model cục bộ bọc metadata job và log xử lý.",
    },
    {
        "method": "GET",
        "path": "/transcripts/{id}",
        "purpose": "Đọc transcript đầy đủ, gồm raw, corrected, segment và speaker alignment.",
        "request": "Path parameter id.",
        "response": "Đối tượng TranscriptView.",
        "errors": "404 transcript không tồn tại; 409 artifact chưa sinh xong.",
        "mapping": "Đọc từ Transcript, SpeakerSegment, AlignmentService, artifact trong output/.",
    },
    {
        "method": "PATCH",
        "path": "/transcripts/{id}/corrections",
        "purpose": "Lưu chỉnh sửa thủ công hoặc yêu cầu correction lại một phần transcript.",
        "request": "JSON gồm segment cần sửa, kiểu sửa manual|phonetic|protonx|llm, và ghi chú điều tra viên.",
        "response": "Transcript sau cập nhật, kèm audit trail của lần chỉnh sửa.",
        "errors": "400 payload sai; 404 transcript không có; 409 transcript đang bị khóa chỉnh sửa.",
        "mapping": "Wrapper mới quanh CorrectionService, VietnamesePhoneticCorrector, ProtonXAdapter và persistence layer cục bộ.",
    },
    {
        "method": "GET",
        "path": "/reports/{id}",
        "purpose": "Lấy báo cáo phân tích forensic có cấu trúc.",
        "request": "Path parameter id.",
        "response": "Đối tượng AnalysisReportView.",
        "errors": "404 report không tồn tại.",
        "mapping": "Đọc artifact JSON trong output/ và ánh xạ về view model GUI.",
    },
    {
        "method": "POST",
        "path": "/reports/{id}/export",
        "purpose": "Xuất gói báo cáo ở dạng TXT, JSON, DOCX hoặc bundle hồ sơ.",
        "request": "JSON gồm định dạng đích, danh sách artifact và metadata hồ sơ cần chèn.",
        "response": "Đường dẫn file xuất và checksum cục bộ.",
        "errors": "404 report không tồn tại; 409 dữ liệu đầu vào thiếu; 500 lỗi export.",
        "mapping": "Wrapper mới quanh OutputFormatter, serializer báo cáo và export service dùng filesystem nội bộ.",
    },
    {
        "method": "GET",
        "path": "/models",
        "purpose": "Liệt kê toàn bộ model offline hiện có trong hệ.",
        "request": "Không có.",
        "response": "Danh sách ModelInfo.",
        "errors": "500 lỗi inventory hoặc quyền truy cập filesystem.",
        "mapping": "Đọc từ core/config.py, SystemFactory và thư mục models/.",
    },
    {
        "method": "POST",
        "path": "/models/{id}/load",
        "purpose": "Chủ động nạp model vào bộ nhớ trước khi chạy job.",
        "request": "Path parameter id, kèm tùy chọn target device nếu cần.",
        "response": "Đối tượng ModelInfo đã cập nhật load_status.",
        "errors": "404 model không có; 409 model chưa offline-ready; 500 load thất bại.",
        "mapping": "Gọi load() trên adapter tương ứng thông qua SystemFactory.",
    },
    {
        "method": "GET",
        "path": "/audit/logs",
        "purpose": "Lấy nhật ký thao tác phục vụ truy vết nghiệp vụ.",
        "request": "Bộ lọc theo job, thời gian, actor hoặc action.",
        "response": "Danh sách AuditEntry.",
        "errors": "500 lỗi đọc log hoặc database cục bộ.",
        "mapping": "Thành phần mới bọc log file, metadata SQLite và action phát sinh từ GUI.",
    },
    {
        "method": "WS",
        "path": "/jobs/{id}/stream",
        "purpose": "Đẩy tiến độ job, log và milestone pipeline theo thời gian thực.",
        "request": "Kết nối websocket theo job_id.",
        "response": "Sự kiện JobStatus, log event và trạng thái artifact mới.",
        "errors": "404 job không tồn tại; 1011 lỗi backend stream.",
        "mapping": "Wrapper streaming quanh orchestration layer mới; không thay đổi adapter lõi.",
    },
]


def load_source_plan_data() -> dict[str, list[str]]:
    data = {"goals": [], "questions": [], "references": []}
    if not SOURCE_PLAN_PATH.exists():
        return data
    from docx import Document as SourceDocument

    source = SourceDocument(SOURCE_PLAN_PATH)
    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("(1)") or text.startswith("(2)") or text.startswith("(3)") or text.startswith("(4)"):
            data["goals"].append(text)
        if text.startswith("- RQ"):
            data["questions"].append(text)
        if text.startswith("[") and "]" in text:
            data["references"].append(text)
    return data


def inventory_counts(path: Path) -> int:
    return sum(1 for file in path.rglob("*") if file.is_file())


def module_inventory() -> list[list[str]]:
    items = [
        ("core", ROOT_DIR / "core", "Entity miền, port interface, cấu hình chung và dịch vụ alignment/output."),
        ("application", ROOT_DIR / "application", "Application service và use case orchestration cho chép lời, correction, phân tích."),
        ("infrastructure", ROOT_DIR / "infrastructure", "Adapter ASR, diarization, correction, LLM, VAD và factory nối lõi với model."),
        ("prompts", ROOT_DIR / "prompts", "Template Jinja2, module prompt, grammar và scenario YAML cho phân tích forensic."),
        ("scripts", ROOT_DIR / "scripts", "Tập script benchmark, setup model, verify pipeline và kiểm tra nhanh từng module."),
        ("tests", ROOT_DIR / "tests", "Test cho prompt, deep analysis và khái quát hóa ngữ cảnh phân tích."),
        ("models", ROOT_DIR / "models", "Kho model offline cho ASR, diarization, correction, VAD và LLM."),
        ("docs", ROOT_DIR / "docs", "Tài liệu nghiên cứu, ghi chú chu trình nâng cấp và báo cáo đánh giá."),
        ("output", ROOT_DIR / "output", "Transcript, diarization, JSON report và artifact chạy thử phục vụ minh chứng."),
    ]
    return [[name, safe_relative(path), str(inventory_counts(path)), role] for name, path, role in items]


def model_inventory() -> list[list[str]]:
    rows = []
    for item in sorted((ROOT_DIR / "models").iterdir()):
        if not item.is_dir():
            continue
        files = [file for file in item.rglob("*") if file.is_file()]
        size = sum(file.stat().st_size for file in files)
        representative = ", ".join(file.name for file in files[:3]) or "Không có"
        rows.append([item.name, human_size(size), str(len(files)), representative[:120], "offline-ready" if files else "trống"])
    return rows


def describe_script(name: str) -> str:
    lowered = name.lower()
    if "benchmark" in lowered:
        return "Benchmark hoặc đo hiệu năng."
    if "setup" in lowered or "download" in lowered:
        return "Chuẩn bị model, dữ liệu hoặc môi trường."
    if "test" in lowered or "verify" in lowered:
        return "Kiểm thử nhanh hoặc xác minh module."
    if "transcribe" in lowered or "diarize" in lowered or "pipeline" in lowered:
        return "Chạy pipeline xử lý âm thanh."
    if "vocab" in lowered or "correction" in lowered:
        return "Kiểm tra hoặc xây dựng thành phần correction hoặc vocabulary."
    return "Tác vụ tiện ích phục vụ nghiên cứu."


def describe_test(name: str) -> str:
    lowered = name.lower()
    if "prompt" in lowered:
        return "Xác minh prompt render hoặc logic prompt."
    if "analysis" in lowered:
        return "Kiểm tra phân tích forensic bằng LLM."
    if "general" in lowered:
        return "Kiểm tra khả năng khái quát hóa hoặc bám scenario."
    return "Kiểm thử chức năng lõi."


def describe_output(name: str) -> str:
    lowered = name.lower()
    if lowered.endswith(".json"):
        return "Artifact JSON phục vụ báo cáo, segment hoặc phân tích."
    if "transcript" in lowered:
        return "Transcript raw hoặc corrected."
    if "diarized" in lowered:
        return "Kết quả gán speaker hoặc transcript theo người nói."
    if lowered.endswith(".txt"):
        return "Bản text phục vụ kiểm tra tay hoặc trình bày."
    return "Artifact minh chứng đầu ra pipeline."


def file_table_rows(folder: Path, description_func) -> list[list[str]]:
    return [[safe_relative(file), description_func(file.name)] for file in sorted(path for path in folder.rglob("*") if path.is_file())]


def implemented_vs_proposed_rows() -> list[list[str]]:
    rows = []
    for item in TECHNIQUE_STATUS:
        rows.append([item["name"], "Đã triển khai", item["status"], "; ".join(item["evidence"][:2])])
    for item in PROPOSALS:
        rows.append([item["name"], "Đề xuất/Nâng cấp", "mới ở mức adapter/probe", "; ".join(item["sources"][:2])])
    return rows


def add_title_page(document: Document) -> None:
    add_center_paragraph(document, "TÀI LIỆU TỔNG HỢP KỸ THUẬT", size=16, bold=True)
    add_center_paragraph(document, "CHERRY CORE 2026", size=18, bold=True)
    document.add_paragraph()
    add_center_paragraph(document, "Hệ lõi trích xuất thông tin điều tra từ dữ liệu âm thanh tiếng Việt", size=14, bold=True)
    add_center_paragraph(document, "Phạm vi: toàn bộ mã nguồn, tài liệu, script, model, prompt và output trong thư mục cherry_core", size=13)
    document.add_paragraph()
    add_center_paragraph(document, "Ngày biên soạn: 13/03/2026", size=13)
    add_center_paragraph(document, "Nguồn tổng hợp: hồ sơ kế hoạch NCKH và khảo sát repo Cherry Core", size=13)
    add_page_break(document)


def add_summary(document: Document, source_plan: dict[str, list[str]]) -> None:
    document.add_heading("Tóm tắt", level=1)
    add_justified_paragraph(document, "Tài liệu này tổng hợp toàn bộ nền tảng kỹ thuật của Cherry Core - lõi xử lý âm thanh điều tra tiếng Việt được xây theo hướng offline-first, module hóa và sẵn sàng mở rộng sang hệ thống giao diện web phục vụ điều tra viên. Tài liệu không viết lại hồ sơ kế hoạch NCKH gốc, mà chuyển hóa kế hoạch đó thành một bức tranh kỹ thuật bám sát hiện trạng repo, qua đó cho thấy những thành phần đã có mã nguồn hoặc minh chứng chạy thử và những nhánh vẫn đang ở mức đề xuất nghiên cứu.")
    add_justified_paragraph(document, "Trong phạm vi cherry_core, các khối kỹ thuật nổi bật gồm: nhận dạng tiếng nói tiếng Việt bằng Whisper V2, PhoWhisper và nhánh WhisperX; tiền xử lý và giảm ảo giác bằng Silero VAD, cấu hình giải mã bảo thủ, Bag-of-Hallucinations và delooping; speaker diarization bằng Pyannote hoặc SpeechBrain ECAPA-TDNN kết hợp VBx; hiệu chỉnh transcript bằng rule-based phonetic correction, ProtonX và LLM cục bộ; cuối cùng là lớp phân tích forensic dựa trên prompt module, scenario YAML và LLM cục bộ llama.cpp hoặc vLLM.")
    if source_plan["goals"]:
        add_justified_paragraph(document, "Các mục tiêu nghiên cứu trong kế hoạch NCKH gốc được tái sử dụng như mốc định hướng kỹ thuật:")
        add_bullets(document, source_plan["goals"][:4])
    add_justified_paragraph(document, "Bên cạnh phần mô tả hệ lõi, tài liệu dành riêng một chương để đặc tả giao diện web hoàn thiện dưới góc nhìn sản phẩm kỹ thuật. Chương này không khóa framework frontend cụ thể, nhưng chuẩn hóa màn hình, DTO, endpoint v1, nguyên tắc lưu trữ offline và cách bọc các use case hiện có của Cherry Core thành API phục vụ GUI.")


def add_abbreviations(document: Document) -> None:
    document.add_heading("Danh mục viết tắt", level=1)
    add_table(document, ["Từ viết tắt", "Giải nghĩa"], [[short, long] for short, long in ABBREVIATIONS])


def add_chapter_1(document: Document, source_plan: dict[str, list[str]]) -> None:
    document.add_page_break()
    document.add_heading("Chương 1. Bối cảnh, mục tiêu và phạm vi khảo sát", level=1)
    add_justified_paragraph(document, "Bài toán cốt lõi của Cherry Core là khai thác thông tin từ âm thanh tiếng Việt trong bối cảnh điều tra và trinh sát kỹ thuật. Không giống các sản phẩm chuyển giọng nói sang văn bản thông thường, một hệ thống phục vụ điều tra phải đồng thời thỏa mãn bốn yêu cầu: độ chính xác ASR đủ cao trên dữ liệu nhiễu; khả năng phân biệt người nói; cơ chế kiểm soát ảo giác của cả ASR lẫn LLM; và vận hành hoàn toàn offline để bảo vệ dữ liệu nhạy cảm.")
    add_justified_paragraph(document, "Hồ sơ kế hoạch NCKH gốc đặt ra các câu hỏi nghiên cứu xoay quanh giảm hallucination, speaker diarization, khai phá thông tin tình báo bằng LLM và kiến trúc pipeline offline. Trong tài liệu này, các câu hỏi đó được diễn giải lại dưới góc nhìn kỹ thuật triển khai của cherry_core, tức là tập trung vào những gì lõi hệ thống đang làm, đang thiếu và cần hoàn thiện.")
    if source_plan["questions"]:
        add_justified_paragraph(document, "Các câu hỏi nghiên cứu chính được kế thừa trực tiếp từ kế hoạch nguồn:")
        add_bullets(document, source_plan["questions"][:4])
    add_justified_paragraph(document, "Phạm vi khảo sát chỉ gồm cherry_core, bao gồm các thư mục core, application, infrastructure, prompts, scripts, tests, models, docs và output. Dự án vietnamese-stt-project cùng các tài liệu cũ ở root workspace bị loại khỏi tài liệu này để tránh pha trộn giữa phần lõi nghiên cứu và một nhánh ứng dụng web khác.")
    add_justified_paragraph(document, "Phương pháp tổng hợp sử dụng bốn bước: đọc tài liệu kế hoạch NCKH do người dùng cung cấp; khảo sát mã nguồn để xác định port, adapter, entity, use case và script; đối chiếu artifact trong output và các báo cáo nghiên cứu trong docs; phân loại từng kỹ thuật theo bốn nhãn trạng thái nhằm giữ tính trung thực giữa nội dung đã triển khai và nội dung mới ở mức đề xuất.")
    document.add_heading("1.1. Nguyên tắc phân loại trạng thái kỹ thuật", level=2)
    add_table(document, ["Nhãn", "Ý nghĩa"], [[label, desc] for label, desc in STATUS_LEGEND])
    document.add_heading("1.2. Kết quả khảo sát tổng quát", level=2)
    add_bullets(document, ["Cherry Core theo kiến trúc Hexagonal, ưu tiên tách domain khỏi adapter model.", "Lõi hiện có nhiều nhánh ASR và diarization song song, cho thấy dự án đang ở trạng thái nghiên cứu so sánh.", "Các thành phần correction và forensic prompt đã phong phú hơn lớp API hoặc UI; giao diện đồ họa chưa được xây ở mức sản phẩm.", "Một số điểm tích hợp còn lệch chuẩn interface, vì vậy tài liệu phải phân biệt rõ phần chạy được với phần còn cần ổn định hóa."])


def add_chapter_2(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Chương 2. Kiến trúc lõi Cherry Core", level=1)
    add_justified_paragraph(document, "Cherry Core được tổ chức theo mô hình Hexagonal Architecture, trong đó phần domain giữ vai trò trung tâm và mọi adapter mô hình AI chỉ đi qua port interface. Cách tổ chức này phù hợp với bối cảnh nghiên cứu vì cho phép thay thế backend ASR, diarization hoặc LLM mà không làm vỡ toàn bộ luồng nghiệp vụ.")
    document.add_heading("2.1. Các lớp kiến trúc", level=2)
    add_table(document, ["Lớp", "Thành phần tiêu biểu", "Vai trò"], [["Domain", "core/domain/entities.py", "Định nghĩa Transcript, SpeakerSegment, StrategicReport."], ["Ports", "core/ports/*.py", "Chuẩn hóa giao tiếp của ASR, diarization, correction và LLM."], ["Application", "application/services, application/use_cases", "Điều phối chép lời, correction, prompt render và phân tích."], ["Infrastructure", "infrastructure/adapters, infrastructure/factories", "Kết nối model thực với port interface, quản lý factory."], ["Presentation", "presentation/cli/main.py", "Điểm vào CLI hiện đại cho người vận hành kỹ thuật."], ["Artifacts", "models/, output/, docs/", "Lưu model offline, đầu ra chạy thử và báo cáo nghiên cứu."]])
    document.add_heading("2.2. Entity và port lõi", level=2)
    add_table(document, ["Nhóm", "Tên", "Vai trò kỹ thuật"], [["Entity", "Transcript", "Lưu text, segment và metadata của ASR."], ["Entity", "SpeakerSegment", "Biểu diễn đoạn lời nói gắn speaker và timestamp."], ["Entity", "StrategicReport", "Báo cáo tình báo chiến lược hoặc tác chiến sau phân tích LLM."], ["Port", "ITranscriber", "Hợp đồng audio -> transcript."], ["Port", "ISpeakerDiarizer", "Hợp đồng audio -> speaker segments."], ["Port", "ITextCorrector", "Hợp đồng hiệu chỉnh transcript."], ["Port", "ILLMEngine", "Hợp đồng nạp model và sinh văn bản cho tác vụ phân tích."]])
    document.add_heading("2.3. Luồng dữ liệu lõi", level=2)
    add_bullets(document, ["Audio đầu vào đi vào ASR adapter thông qua ITranscriber.", "Nếu bật correction, TranscribeAudioUseCase chạy thêm ITextCorrector hoặc service correction nhiều tầng.", "Nếu cần speaker-aware transcript, audio đi qua ISpeakerDiarizer, sau đó AlignmentService căn chỉnh speaker với transcript.", "PromptManager nạp template Jinja2 và scenario YAML để dựng prompt forensic.", "LLM adapter sinh báo cáo, sau đó ứng dụng hoặc CLI serialize báo cáo về JSON và text file trong output.", "Toàn bộ model được nạp từ models theo cấu hình cục bộ, không phụ thuộc cloud trong pha vận hành."])
    document.add_heading("2.4. Quản lý model offline và artifact", level=2)
    add_justified_paragraph(document, "Thư mục models chứa đầy đủ model cho ASR, correction, diarization, VAD và LLM. Điều này thể hiện nguyên tắc offline-first xuyên suốt: PhoWhisper, Whisper V2/V3, SpeechBrain, ProtonX, Silero, Vistral, Qwen3 và nhánh WhisperX đều có vị trí lưu cục bộ. Song song đó, output đóng vai trò evidence store, nơi lưu transcript thô, transcript đã hiệu chỉnh, speaker transcript và JSON report để phục vụ đánh giá kỹ thuật lẫn nghiệp vụ.")
    document.add_heading("2.5. Điểm vào hiện tại của hệ", level=2)
    add_bullets(document, ["presentation/cli/main.py: entry point tương đối hiện đại, dùng TranscribeAudioUseCase và GenerateStrategicReportUseCase.", "cli.py: entry point cũ hơn, cho thấy dấu vết refactor và một số import hoặc service không còn đồng bộ.", "Các script trong scripts: công cụ nghiên cứu, benchmark, setup và verify từng mô-đun."])


def add_chapter_3(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Chương 3. Kỹ thuật đã triển khai hoặc đã có minh chứng trong Cherry Core", level=1)
    add_justified_paragraph(document, "Chương này chỉ tập trung vào các kỹ thuật đã hiện diện trong mã nguồn hoặc đã có dấu vết kiểm thử, output, script probe trong repo. Mỗi kỹ thuật được gắn duy nhất một nhãn trạng thái nhằm tránh trình bày quá mức mức độ hoàn thiện.")
    add_table(document, ["Kỹ thuật", "Nhãn trạng thái", "Tóm tắt", "Minh chứng tiêu biểu"], [[item["name"], item["status"], item["summary"], '; '.join(item["evidence"][:2])] for item in TECHNIQUE_STATUS])
    document.add_heading("3.1. Cụm kỹ thuật ASR", level=2)
    add_justified_paragraph(document, "Nhánh ASR của Cherry Core không đi theo một engine duy nhất. SystemFactory hiện chọn Whisper V2 làm champion mặc định để ưu tiên độ ổn định và giảm hallucination, trong khi PhoWhisper được giữ như một lựa chọn tăng cường tiếng Việt và WhisperX được phát triển như nhánh hợp nhất ASR + alignment + diarization.")
    add_bullets(document, ["Whisper V2: adapter chạy openai-whisper offline, có hậu xử lý chống lặp và có output minh chứng trong output.", "PhoWhisper: ưu tiên model local dạng safetensors, thích hợp cho tiếng Việt nhưng hiện mới chứng minh chắc ở mức code và model local.", "WhisperX: thể hiện hướng nâng cấp mạnh nhất cho GUI timeline người nói, nhưng đang ở trạng thái kiểm chứng pipeline offline."])
    document.add_heading("3.2. Giảm ảo giác trong ASR", level=2)
    add_justified_paragraph(document, "Cherry Core dùng chiến lược giảm ảo giác nhiều lớp. Lớp thứ nhất là cấu hình giải mã bảo thủ trong WhisperV2Adapter; lớp thứ hai là Silero VAD loại bỏ khoảng lặng; lớp thứ ba là HallucinationFilter dùng Bag-of-Hallucinations và delooping. Cách tiếp cận này bám sát mục tiêu nghiên cứu trong kế hoạch NCKH: giảm bịa nội dung khi âm thanh kém chất lượng hoặc có vùng im lặng.")
    add_bullets(document, ["Cấu hình condition_on_previous_text=False cắt chuỗi phụ thuộc lịch sử nhằm tránh lây lan lỗi.", "Ngưỡng compression_ratio_threshold, logprob_threshold, no_speech_threshold chặn segment đáng ngờ.", "Silero VAD có ngưỡng bảo thủ để ưu tiên không bỏ sót tiếng nói trong bối cảnh điều tra.", "Bag-of-Hallucinations mã hóa cả mẫu tiếng Việt lẫn tiếng Anh, phù hợp thực tế audio lẫn tạp âm hoặc slogan chèn.", "Delooping xử lý vòng lặp từ và cụm, điển hình trong các trường hợp Whisper tự lặp nhiều lần tên riêng hoặc filler."])
    document.add_heading("3.3. Speaker diarization và căn chỉnh người nói", level=2)
    add_justified_paragraph(document, "Nhánh diarization có tính chất nghiên cứu so sánh rõ rệt. Repo giữ song song Pyannote, SpeechBrain ECAPA-TDNN, EnhancedDiarizer dựa trên Resemblyzer và lớp VBx làm mượt nhãn speaker. Phần căn chỉnh AlignmentService đóng vai trò liên kết tất cả nhánh này về một transcript theo người nói.")
    add_bullets(document, ["PyannoteAdapter hướng tới SOTA nhưng hiện vẫn chịu ràng buộc token hoặc license.", "SpeechBrainAdapter là lựa chọn offline rõ ràng hơn, có sẵn model cục bộ và kết nối trực tiếp với VBxRefiner.", "VBxRefiner dùng logic HMM hoặc Viterbi để giảm nhảy speaker không tự nhiên.", "AlignmentService dùng IntervalTree khi có thư viện, và linear fallback khi thiếu phụ thuộc."])
    document.add_heading("3.4. Hậu hiệu chỉnh tiếng Việt", level=2)
    add_justified_paragraph(document, "Lớp correction của Cherry Core đi theo cấu trúc nhiều tầng. Tầng một là VietnamesePhoneticCorrector sửa các lỗi đồng âm, lỗi phát âm và tên riêng dựa trên từ điển JSON. Tầng hai là ProtonXAdapter, một mô hình Seq2Seq cục bộ tối ưu cho correction văn bản dài. Tầng ba là CorrectionService kết hợp vocabulary injection và LLM cục bộ để sửa theo ngữ cảnh miền.")
    add_bullets(document, ["Assets vocabulary cho thấy hệ được chuẩn bị cho các miền chuyên biệt như an ninh, pháp lý, khách sạn, tiếng lóng.", "PROOF_RESULTS.md chứng minh correction đã sửa được một số lỗi homophone và entity đáng chú ý.", "Thiết kế hiện tại phù hợp để GUI cho phép chạy correction tự động rồi cho điều tra viên duyệt lại thủ công."])
    document.add_heading("3.5. Phân tích forensic bằng LLM", level=2)
    add_justified_paragraph(document, "Mảng forensic LLM là nơi Cherry Core thể hiện tham vọng vượt khỏi hệ ASR thông thường. PromptManager dùng Jinja2 để ghép module 5W1H, SVA, SCAN, cảm xúc, nhạy cảm, suy luận vai trò và scenario YAML thành prompt điều tra. AnalysisService hiện dùng LlamaCppAdapter làm engine mặc định, đồng thời repo đã chuẩn bị VLLMAdapter cho nhánh throughput cao hơn.")
    add_bullets(document, ["Prompt deep_investigation.j2 ép đầu ra JSON-only, nhằm hướng tới báo cáo forensic có cấu trúc.", "Scenario YAML cho thấy hệ nhắm vào nhiều bối cảnh nghiệp vụ như ma túy, tình báo chung và lừa đảo công nghệ cao.", "Artifact output/..._analysis_report.json minh chứng lớp LLM đã được nối vào pipeline chạy thật.", "Tuy vậy, sample output hiện còn lưu cả raw_output với lỗi parse JSON, cho thấy structured generation chưa ổn định hoàn toàn."])
    document.add_heading("3.6. Bằng chứng chạy thử và công cụ nghiên cứu", level=2)
    add_justified_paragraph(document, "Ngoài mã nguồn, repo có hệ script và output khá dày. Đây là lớp evidence quan trọng để phân biệt một ý tưởng nghiên cứu với một thực thể kỹ thuật đã từng được kiểm chứng. scripts chứa benchmark, verify, setup và debug; tests chứa test logic prompt; output giữ transcript, segment JSON và report phân tích.")


def add_chapter_4(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Chương 4. Các hướng nghiên cứu và nâng cấp đã xuất hiện trong repo", level=1)
    add_justified_paragraph(document, "Khối docs của Cherry Core lưu nhiều báo cáo nghiên cứu có giá trị chiến lược. Những tài liệu này không nên bị hiểu nhầm là tính năng đã hoàn thiện; chúng là bản đồ nâng cấp của lõi hệ thống. Bảng dưới đây tổng hợp các nhánh nổi bật nhất cùng lợi ích, khoảng trống hiện tại và rủi ro triển khai.")
    add_table(document, ["Kỹ thuật đề xuất", "Lợi ích kỳ vọng", "Khoảng trống hiện tại", "Rủi ro"], [[item["name"], item["benefit"], item["gap"], item["risk"]] for item in PROPOSALS])
    add_justified_paragraph(document, "Tựu trung, định hướng nâng cấp của repo hội tụ vào ba trục: giảm ảo giác sâu hơn ở cả ASR và LLM; nâng chất lượng speaker diarization và structured output; chuẩn bị một hạ tầng LLM mạnh hơn để phục vụ GUI và xử lý nhiều hồ sơ. Đây là các trục hợp lý với đề bài NCKH ban đầu, nhưng cần được trình bày như roadmap có kiểm soát thay vì tính năng đã sẵn sàng đưa vào sản phẩm.")


def add_chapter_5(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Chương 5. Kế hoạch hoàn thiện GUI web cho Cherry Core", level=1)
    add_justified_paragraph(document, "GUI web được đề xuất như một lớp bao ngoài cho cherry_core, phục vụ điều tra viên hoặc nhà phân tích vận hành trên máy nội bộ. Mục tiêu của GUI không phải thay thế lõi xử lý hiện tại, mà biến lõi đó thành một công cụ tác nghiệp có thể nạp audio, theo dõi job, duyệt transcript, chạy phân tích LLM và xuất hồ sơ.")
    document.add_heading("5.1. Nguyên tắc thiết kế", level=2)
    add_bullets(document, ["Offline-first: toàn bộ audio, model, metadata và artifact nằm trong môi trường cục bộ sau khi cài đặt.", "Evidence-oriented: mọi chỉnh sửa, export và phân tích đều phải có audit trail.", "Human-in-the-loop: transcript và báo cáo LLM đều cho phép điều tra viên duyệt lại trước khi chốt.", "Frontend-neutral: đặc tả API và DTO độc lập với React, HTMX hay framework cụ thể.", "Không multi-tenant ở v1: tập trung một máy, một đơn vị vận hành, không thiết kế SSO hoặc cộng tác thời gian thực nhiều người."])
    document.add_heading("5.2. Các màn hình bắt buộc", level=2)
    add_table(document, ["Màn hình", "Mục tiêu", "Dữ liệu vào", "Dữ liệu ra", "Ánh xạ lõi"], [[item["name"], item["goal"], item["inputs"], item["outputs"], item["core_mapping"]] for item in GUI_SCREENS])
    document.add_heading("5.3. Luồng nghiệp vụ chuẩn", level=2)
    add_bullets(document, ["Bước 1: người dùng tạo hồ sơ hoặc chọn hồ sơ sẵn có trên màn hình tổng quan.", "Bước 2: tải file audio và chọn engine ASR, bật correction hoặc diarization khi cần.", "Bước 3: hệ tạo JobStatus và khởi động pipeline chép lời.", "Bước 4: khi transcript hoàn thành, người dùng xem raw text, corrected text và timeline người nói.", "Bước 5: nếu có sai sót, người dùng chỉnh thủ công hoặc yêu cầu correction lại một phần transcript.", "Bước 6: transcript được duyệt đưa sang panel forensic để chạy prompt theo scenario.", "Bước 7: báo cáo phân tích được đọc, kiểm tra mức đe dọa, entity, timeline và khuyến nghị.", "Bước 8: người dùng xuất gói TXT, JSON, DOCX và hệ lưu audit trail cho toàn bộ thao tác."])
    document.add_heading("5.4. Contract dữ liệu công khai", level=2)
    dto_index = 1
    for dto_name, fields in DTO_DEFINITIONS.items():
        document.add_heading(f"5.4.{dto_index}. {dto_name}", level=3)
        add_table(document, ["Trường", "Kiểu", "Ý nghĩa"], [list(field) for field in fields])
        dto_index += 1
    document.add_heading("5.5. Bộ endpoint v1", level=2)
    add_table(document, ["Method", "Path", "Mục đích", "Đầu vào", "Đầu ra", "Lỗi thường gặp", "Ánh xạ về Cherry Core"], [[item["method"], item["path"], item["purpose"], item["request"], item["response"], item["errors"], item["mapping"]] for item in ENDPOINT_SPECS])
    document.add_heading("5.6. Thiết kế lưu trữ và đóng gói offline", level=2)
    add_bullets(document, ["Metadata job, transcript và audit trail dùng SQLite cục bộ để thuận tiện backup, tra cứu và triển khai một máy.", "Audio và artifact lớn tiếp tục lưu trên filesystem nội bộ theo thư mục hồ sơ; DB chỉ giữ metadata và checksum.", "Model inventory đọc trực tiếp từ models, tránh đồng bộ kép.", "Dịch vụ GUI có thể tách thành API wrapper, background worker và frontend bundle tĩnh; tuy nhiên v1 vẫn có thể chạy trong một tiến trình ứng dụng duy nhất.", "Mọi export DOCX phải đóng dấu nguồn transcript, model đã dùng, thời gian xử lý và cảnh báo về mức độ kiểm chứng của báo cáo LLM."])


def add_chapter_6(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Chương 6. Hướng hoàn thiện, tiêu chí đánh giá và lộ trình nâng cấp", level=1)
    add_justified_paragraph(document, "Cherry Core đã có một lõi kỹ thuật đáng kể, nhưng để trở thành hệ thống nghiên cứu - triển khai đủ vững cho môi trường điều tra thì cần thêm một pha ổn định hóa. Pha này không chỉ là bổ sung tính năng, mà còn là chuẩn hóa interface, khóa pipeline mặc định và làm rõ ranh giới giữa nhánh legacy với nhánh nghiên cứu.")
    document.add_heading("6.1. Các điểm mạnh hiện tại", level=2)
    add_bullets(document, ["Có nhiều lựa chọn ASR, diarization và LLM để so sánh trong cùng một lõi kiến trúc.", "Các thành phần offline asset đã hiện diện đầy đủ trong models, giảm phụ thuộc runtime download.", "Prompt forensic được xây tương đối bài bản, có scenario và module riêng cho nhiều khía cạnh phân tích.", "Repo giữ lại được output và báo cáo nghiên cứu, thuận tiện viết báo cáo NCKH và thiết kế GUI tiếp theo."])
    document.add_heading("6.2. Khoảng trống kỹ thuật cần khóa trước khi bọc GUI", level=2)
    add_bullets(document, ["Cần thống nhất interface LLM: GenerateStrategicReportUseCase gọi analyze() trong khi ILLMEngine công khai generate() và load().", "Cần đồng bộ entity hoặc domain với adapter: ví dụ PhoWhisperAdapter truyền thêm trường language không có trong dataclass hiện tại.", "Cần tách rõ entry point cũ và mới: cli.py còn dấu vết import hoặc service cũ, trong khi presentation/cli/main.py là nhánh hiện đại hơn.", "Structured JSON output chưa đủ ổn định; sample report còn phải lưu raw_output do parse lỗi.", "Chưa có persistence layer và job orchestration chính thức cho GUI; hiện repo vẫn thiên về script và CLI."])
    document.add_heading("6.3. Tiêu chí đánh giá thành công cho giai đoạn hoàn thiện", level=2)
    add_bullets(document, ["ASR: giữ hoặc cải thiện WER mục tiêu so với baseline hiện tại trên tập tiếng Việt thực tế.", "Hallucination: giảm rõ rệt segment lặp hoặc bịa ở vùng silence và có cơ chế flag transcript đáng ngờ.", "Diarization: có benchmark DER nhất quán cho ít nhất một pipeline mặc định.", "LLM: valid rate của JSON đạt mức chấp nhận được để GUI không phải dựa vào raw_output làm nguồn chính.", "Product: GUI MVP đủ các màn hình và endpoint đã nêu, hỗ trợ quy trình human-in-the-loop trọn vẹn."])
    document.add_heading("6.4. Lộ trình triển khai đề xuất", level=2)
    add_bullets(document, ["Giai đoạn 1 - Ổn định lõi: khóa interface, chọn pipeline mặc định, chuẩn hóa artifact và sửa các điểm mismatch port hoặc entity.", "Giai đoạn 2 - Tăng độ tin cậy: hoàn thiện structured output, bổ sung benchmark DER/WER hoặc hallucination và đóng gói audit trail.", "Giai đoạn 3 - Bọc API nội bộ: tạo persistence SQLite, background jobs và read-model cho transcript hoặc report.", "Giai đoạn 4 - GUI MVP: hiện thực các màn hình đã đặc tả, đưa transcript hoặc speaker timeline vào thao tác thực tế.", "Giai đoạn 5 - Nâng cấp nghiên cứu: CoVe, Calm-Whisper, XGrammar, semantic retrieval, vLLM serving ổn định."])
    document.add_heading("Kết luận", level=2)
    add_justified_paragraph(document, "Ở thời điểm khảo sát tháng 03/2026, cherry_core đã vượt qua mức một prototype đơn lẻ và hình thành một lõi kỹ thuật khá đầy đủ cho bài toán điều tra âm thanh tiếng Việt. Giá trị lớn nhất của repo nằm ở chỗ nó gom được nhiều hướng nghiên cứu về ASR, diarization, correction và forensic LLM trong cùng một khung kiến trúc. Giá trị còn lại cần hoàn thiện là biến khung đó thành sản phẩm tác nghiệp có giao diện, persistence và các chuẩn dữ liệu ổn định.")


def add_appendices(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Phụ lục A. Bảng inventory mô-đun", level=1)
    add_table(document, ["Nhóm", "Đường dẫn", "Số file", "Vai trò"], module_inventory())
    document.add_heading("Phụ lục B. Bảng model offline", level=1)
    add_table(document, ["Model folder", "Dung lượng", "Số file", "Tệp tiêu biểu", "Trạng thái"], model_inventory())
    document.add_heading("Phụ lục C. Bảng script minh chứng", level=1)
    add_table(document, ["Script", "Vai trò"], file_table_rows(ROOT_DIR / "scripts", describe_script))
    document.add_heading("Phụ lục D. Bảng test minh chứng", level=1)
    add_table(document, ["Test file", "Vai trò"], file_table_rows(ROOT_DIR / "tests", describe_test))
    document.add_heading("Phụ lục E. Bảng output minh chứng", level=1)
    add_table(document, ["Artifact", "Vai trò"], file_table_rows(ROOT_DIR / "output", describe_output))
    document.add_heading("Phụ lục F. Mapping implemented vs proposed", level=1)
    add_table(document, ["Hạng mục", "Nhóm", "Nhãn", "Minh chứng"], implemented_vs_proposed_rows())
    document.add_heading("Phụ lục G. Glossary thuật ngữ", level=1)
    add_table(document, ["Thuật ngữ", "Giải nghĩa"], [[term, meaning] for term, meaning in GLOSSARY])
    document.add_heading("Phụ lục H. Tài liệu tham khảo", level=1)
    for reference in REFERENCES:
        add_justified_paragraph(document, reference)


def add_source_reference_note(document: Document, source_plan: dict[str, list[str]]) -> None:
    if not source_plan["references"]:
        return
    document.add_heading("Ghi chú nguồn từ kế hoạch NCKH", level=1)
    add_justified_paragraph(document, "Các trích dẫn học thuật gốc xuất hiện trong kế hoạch NCKH của đề tài được giữ lại như nguồn định hướng cho phần nghiên cứu và nâng cấp của Cherry Core.")
    for reference in source_plan["references"][:10]:
        add_justified_paragraph(document, reference)


def build_document() -> Document:
    source_plan = load_source_plan_data()
    document = Document()
    configure_document(document)
    add_title_page(document)
    add_summary(document, source_plan)
    add_abbreviations(document)
    add_chapter_1(document, source_plan)
    add_chapter_2(document)
    add_chapter_3(document)
    add_chapter_4(document)
    add_chapter_5(document)
    add_chapter_6(document)
    add_appendices(document)
    add_source_reference_note(document, source_plan)
    return document


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
